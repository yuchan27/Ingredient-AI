from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "C11215139_medium_Flutter.docx"
OUTPUT = ROOT / "C11215139_medium_Flutter_FoodLens_AI_final.docx"
DOCS = ROOT / "foodlens_ai" / "docs"
SCREENSHOTS = DOCS / "screenshots"
SAMPLES = DOCS / "assets" / "ai_samples"

FONT = "Microsoft JhengHei"
GREEN = "13795B"
LIGHT_GREEN = "E7F4ED"
LIGHT_GRAY = "F3F5F3"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, 9.5, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, 18 if level == 1 else 13, bold=True, color=GREEN if level == 1 else None)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, 10.5, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, 10.5)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 10.5)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"• {text}")
    set_run_font(run, 10)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_properties)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index], GREEN)
        if widths:
            table.rows[0].cells[index].width = Inches(widths[index])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], value, align=align)
            if widths:
                cells[index].width = Inches(widths[index])
            if row_index % 2 == 1:
                shade_cell(cells[index], LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_image_pair(doc, left_path, left_caption, right_path, right_caption, width=3.0):
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    for column, (path, caption) in enumerate(((left_path, left_caption), (right_path, right_caption))):
        image_paragraph = table.rows[0].cells[column].paragraphs[0]
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(str(path), width=Inches(width))
        caption_paragraph = table.rows[1].cells[column].paragraphs[0]
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.add_run(caption)
        set_run_font(run, 9, bold=True, color=GREEN)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_report():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(SOURCE)
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("FoodLens AI MVP 完整擴充實作")
    set_run_font(run, 24, bold=True, color=GREEN)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("以原 medium Flutter 期末專案為基礎的雲端、AI 與飲食分析升級")
    set_run_font(run, 12, color="4E5A53")
    add_body(doc, "本擴充章節保留前述 medium 報告的原始內容與圖片，並在其後補充實際完成的 FoodLens AI Flutter App、Node.js API、Firebase 資料架構、Gemini 圖片分析與 Android Emulator 驗證。")

    add_heading(doc, "1. 升級目標與系統定位")
    add_body(doc, "FoodLens AI 將原先的食品成分分析概念擴充為可長期記錄的飲食分析 App。使用者可以用 Email 註冊、上傳圖片、讓 AI 估算熱量與三大營養素，再依實際份量手動修正，最後寫入 Firestore。")
    add_table(
        doc,
        ["層級", "技術", "主要責任"],
        [
            ["Android App", "Flutter / Material 3", "帳號、圖片選擇、營養編輯、紀錄與分析"],
            ["身分與資料", "Firebase Auth / Firestore / Storage", "Email 驗證、離線快取、使用者私有資料與圖片"],
            ["API", "Express 5 / Firebase Admin", "驗證 ID Token、檢查 Storage path 所有權、回傳固定 JSON"],
            ["AI", "Gemini 2.5 Flash / Gemma 4 fallback", "圖片食品辨識、熱量與營養估算"],
            ["郵件", "Firebase verification / Nodemailer", "註冊驗證信與開發環境 SMTP 健康檢查"],
        ],
        widths=[1.15, 1.7, 3.35],
    )

    add_heading(doc, "2. 已實作功能")
    for item in [
        "Email/Password 註冊、登入、登出，並在驗證 Email 前限制進入資料頁。",
        "相機或圖庫選圖，上傳 Firebase Storage 後由受 Firebase ID Token 保護的 API 進行分析。",
        "AI 結果包含食品名、熱量、蛋白質、脂肪、碳水、信心度與備註，儲存前可修改。",
        "依日期與餐別記錄，並加入餐點花費，可從詳細頁修改或刪除。",
        "今日熱量、三大營養素、花費、七日概覽與下一餐建議。",
        "Firestore Android 本機持久快取：離線時可讀取已快取資料，寫入會在聯網後自動同步。",
        "設定頁顯示帳號、API 主機、同步說明、隱私邊界與可儲存至使用者子集合的意見表單。",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "3. Android Emulator 實際執行畫面")
    add_body(doc, "以 Pixel 9 Pro XL AVD（1344×2992）安裝 x86_64 release APK，並使用 ADB UI tree 定位操作。下列截圖皆為實際 Emulator 畫面，不是設計稿。")
    add_image_pair(doc, SCREENSHOTS / "home.png", "圖 1  今日熱量、營養與記帳首頁", SCREENSHOTS / "insights.png", "圖 2  飲食總分析與下一餐建議")
    add_image_pair(doc, SCREENSHOTS / "image-selected.png", "圖 3  從 Android Photo Picker 選擇圖片", SCREENSHOTS / "analysis-result.png", "圖 4  AI 結果可手動修改後儲存")
    doc.add_page_break()
    add_image_pair(doc, SCREENSHOTS / "record-detail.png", "圖 5  紀錄詳細、修改與刪除", SCREENSHOTS / "settings.png", "圖 6  帳號、離線同步與意見表單")

    add_heading(doc, "4. 實際圖片辨識測試")
    add_body(doc, "後端直接使用 Google AI Studio API key 對四張不同排版的中文食品標示圖進行線上測試。四張皆成功解析為固定 JSON，後端會先使用 gemini-2.5-flash，失敗或 JSON 錯誤才切換 gemma-4-26b-a4b-it。")
    add_image_pair(doc, SAMPLES / "chicken-curry-label.png", "樣本 A  雞胸咖哩調理包", SAMPLES / "pudding-label.png", "樣本 B  焦糖布丁營養標示", width=2.65)
    add_image_pair(doc, SAMPLES / "cookie-label.png", "樣本 C  成分與營養教學圖", SAMPLES / "walnut-cake-label.png", "樣本 D  核桃糕食品標示", width=2.65)
    add_table(
        doc,
        ["樣本", "辨識食品", "kcal", "P", "F", "C", "conf."],
        [
            ["A", "雞胸肉咖哩風味調理包", "282", "16.9", "11.6", "27.6", "0.95"],
            ["B", "北歐先生手工焦糖烤布丁", "787", "33", "36", "95.5", "0.90"],
            ["C", "餅乾", "628", "12.4", "24", "90", "1.00"],
            ["D", "擎寶牌-尚好呷核桃糕", "1856", "23.6", "87.6", "243.2", "1.00"],
        ],
        widths=[0.45, 2.45, 0.65, 0.55, 0.55, 0.65, 0.65],
    )
    add_body(doc, "說明：P/F/C 分別代表蛋白質、脂肪與碳水化合物，單位為公克。AI 結果屬估算或標示文字讀取，不取代專業醫療或營養建議。")

    add_heading(doc, "5. API、資料庫與安全設計")
    add_table(
        doc,
        ["API", "保護", "用途"],
        [
            ["GET /health", "公開", "本機與 Cloud Run 健康檢查"],
            ["POST /sendTestEmail", "僅 development", "寄到設定 SMTP_USER，防止郵件轉送濫用"],
            ["POST /analyzeFoodImage", "Firebase Bearer token", "只允許 users/{uid}/ 所有圖片，下載後呼叫 AI"],
        ],
        widths=[1.75, 1.45, 3.45],
    )
    add_body(doc, "Firestore schema：", bold_prefix="Firestore schema：")
    add_bullet(doc, "users/{uid}：使用者主資料。")
    add_bullet(doc, "users/{uid}/food_records/{recordId}：營養、日期、餐別、花費與圖片路徑。")
    add_bullet(doc, "users/{uid}/feedback/{feedbackId}：意見表單。")
    add_body(doc, "Storage path： users/{uid}/food_images/{recordId}.jpg。Firestore Rules 與 Storage Rules 均要求 request.auth.uid == uid，且 Storage 只允許小於 10 MB 的 JPEG/PNG/WebP。")

    add_heading(doc, "6. 自動化與實機驗證結果")
    add_table(
        doc,
        ["驗證項目", "結果", "證據"],
        [
            ["Node API tests", "PASS", "18 tests：auth、path ownership、AI fallback、SMTP 等"],
            ["Flutter tests", "PASS", "10 tests：model、API URL、nutrition summary、HTTP、widget"],
            ["flutter analyze", "PASS", "No issues found"],
            ["Local /health", "PASS", "{ ok: true }"],
            ["Nodemailer", "PASS", "development test email 成功寄出"],
            ["Gemini 圖片", "PASS", "4/4 圖片成功回傳固定 JSON"],
            ["Android APK", "PASS", "debug universal 與 release split-per-ABI 建置成功"],
            ["Pixel 9 Emulator", "PASS", "安裝、啟動、導覽、選圖、新增、儲存、詳細頁"],
        ],
        widths=[1.65, 0.8, 4.25],
    )

    add_heading(doc, "7. 執行方式")
    add_body(doc, "後端（PowerShell）：", bold_prefix="後端（PowerShell）：")
    add_body(doc, "cd C:\\Users\\wuwu6\\StudioProjects\\App_medium\\foodlens_ai\\server\nnpm install\nnpm start")
    add_body(doc, "Android Studio：開啟 foodlens_ai/app，選擇 Pixel 9 Emulator。Firebase 完成前可在 Run configuration 加入 --dart-define=DEMO_MODE=true；正式測試改用 --dart-define-from-file=firebase.dev.json。Android Emulator 的本機 API URL 為 http://10.0.2.2:3000。")
    add_body(doc, "Cloud Run：依 foodlens_ai/docs/deploy-cloud-run.md 建立 Secret Manager secrets，再使用 gcloud run deploy --source . 部署。實機必須將 API_BASE_URL 改為 Cloud Run 或 ngrok HTTPS URL。")

    add_heading(doc, "8. 期末作業連結與交付")
    github = doc.add_paragraph()
    github.paragraph_format.space_after = Pt(6)
    prefix = github.add_run("GitHub 程式碼：")
    set_run_font(prefix, 10.5, bold=True)
    add_hyperlink(github, "Ingredient-AI / codex/foodlens-ai-mvp", "https://github.com/yuchan27/Ingredient-AI/tree/codex/foodlens-ai-mvp")
    video = doc.add_paragraph()
    video.paragraph_format.space_after = Pt(6)
    prefix = video.add_run("執行畫面影片：")
    set_run_font(prefix, 10.5, bold=True)
    add_hyperlink(video, "FoodLens AI Android Emulator demo", "https://github.com/yuchan27/Ingredient-AI/raw/refs/heads/codex/foodlens-ai-mvp/foodlens_ai/docs/demo/foodlens-ai-demo.mp4")
    add_body(doc, "程式、規則、測試、Dockerfile、部署文件、截圖與影片均位於 foodlens_ai/ 目錄。")

    add_heading(doc, "9. 現階段邊界與上線前準備")
    add_body(doc, "本次已完成程式、安全規則、本機 API、SMTP、AI 實際圖片與 Emulator 驗證。但 server/.env 目前尚未有 Firebase Project ID/Storage bucket，因此真實 Firebase 註冊、Firestore 與 Storage 尚未能在此環境做線上整合測試。上線前需完成：")
    for item in [
        "建立 Firebase 專案、啟用 Email/Password、部署 firestore.rules 與 storage.rules。",
        "將 Gemini/SMTP secret 放入 Google Secret Manager，部署 Cloud Run 後更新 API_BASE_URL。",
        "在 Google Play Console 準備隱私政策、資料安全表單、正式 app icon、簽章與帳號/資料刪除流程。",
        "將營養分析標示為估計值，並加入醫療與過敏風險免責說明。",
    ]:
        add_bullet(doc, item)

    for paragraph in doc.paragraphs[-100:]:
        for run in paragraph.runs:
            if run.font.name is None:
                set_run_font(run, 10.5)

    doc.core_properties.title = "FoodLens AI MVP 期末專案擴充報告"
    doc.core_properties.subject = "Flutter, Firebase, Gemini, Cloud Run"
    doc.core_properties.keywords = "FoodLens AI, Flutter, Firebase, Gemini, Android"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
